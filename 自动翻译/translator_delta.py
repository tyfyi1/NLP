#!/usr/bin/env python3
"""
平衡版DelTA多级记忆翻译智能体【PDF+md+docx+txt全格式完整版】
新增PDF解析：自动提取论文文本、识别数学公式并包裹$ $$，中转md完成翻译
适配ICLR/arXiv LaTeX论文，公式完整保护，解决长文本截断
依赖新增：pymupdf (fitz)
安装全部依赖：pip install anthropic sacrebleu bert-score python-docx pymupdf
"""
# 全局环境变量修复冲突 & 强制CPU运行BERT打分
import os
import shutil
import tempfile
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["CUDA_VISIBLE_DEVICES"] = "-1"
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
import json
import time
import random
import argparse
import re
from pathlib import Path
from typing import Optional, Dict, List, Tuple
# 第三方依赖校验
try:
    import anthropic
    from anthropic import RateLimitError, APIStatusError
except ImportError:
    print("缺失anthropic，执行：pip install anthropic")
    import sys
    sys.exit(1)
try:
    import sacrebleu
except ImportError:
    print("缺失sacrebleu，执行：pip install sacrebleu")
    import sys
    sys.exit(1)
try:
    from bert_score import score
except ImportError:
    print("缺失bert-score，执行：pip install bert-score")
    import sys
    sys.exit(1)
# Word文档依赖
try:
    from docx import Document
except ImportError:
    print("如需翻译docx，请执行：pip install python-docx")
    Document = None
# PDF解析依赖
try:
    import fitz
except ImportError:
    print("PDF功能缺失pymupdf，执行：pip install pymupdf")
    fitz = None

# ===================== PDF解析工具：提取文本+自动标记LaTeX公式 =====================
def pdf_to_temp_md(pdf_path: str, temp_dir: str) -> str:
    """
    读取PDF，识别行内/行间公式，自动加$ $$，输出临时md文件路径
    仅支持可复制文本PDF，扫描图片PDF公式失效
    """
    if fitz is None:
        raise RuntimeError("未安装pymupdf，无法解析PDF: pip install pymupdf")
    doc = fitz.open(pdf_path)
    md_lines = []
    for page in doc:
        blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
        for blk in blocks:
            if blk["type"] != 0:
                continue
            line_texts = []
            for line in blk["lines"]:
                line_str = ""
                for span in line["spans"]:
                    text = span["text"].strip()
                    # 简易数学识别：含数学符号判定为行内公式
                    math_chars = {"∑", "∫", "∂", "α", "β", "γ", "δ", "λ", "μ", "σ", "ω", "=", "+", "-", "/", "^", "_"}
                    has_math = any(c in text for c in math_chars)
                    if len(text) <= 8 and has_math and not text.isalnum():
                        line_str += f"${text}$ "
                    else:
                        line_str += text + " "
                line_texts.append(line_str.strip())
            full_block = " ".join(line_texts).strip()
            if not full_block:
                continue
            # 多行大块数学判定为块公式
            if len(full_block) < 30 and any(c in full_block for c in {"∑", "∫", "="}) and full_block.count("=") >= 1:
                md_lines.append(f"\n$$\n{full_block}\n$$\n")
            else:
                md_lines.append(full_block)
                md_lines.append("")
    doc.close()
    temp_md_path = Path(temp_dir) / (Path(pdf_path).stem + "_temp_extract.md")
    temp_md_path.write_text("\n".join(md_lines), encoding="utf-8")
    return str(temp_md_path)

# ===================== 文本分段工具（论文优化：1200字符防截断） =====================
def split_long_text(text: str, max_chunk_chars: int = 1200) -> List[str]:
    paragraphs = re.split(r"\n{2,}", text)
    chunks = []
    current_chunk = []
    current_len = 0
    for para in paragraphs:
        para_strip = para.strip()
        if not para_strip:
            continue
        para_len = len(para_strip)
        if current_len + para_len > max_chunk_chars and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = []
            current_len = 0
        current_chunk.append(para_strip)
        current_len += para_len
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))
    return chunks

# ===================== Word Docx 读写工具 =====================
class DocxHandler:
    @staticmethod
    def read_docx(file_path: str) -> str:
        if Document is None:
            raise RuntimeError("未安装python-docx，执行 pip install python-docx")
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_txt = cell.text.strip()
                    if cell_txt:
                        row_text.append(cell_txt)
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n\n".join(full_text)

    @staticmethod
    def write_docx(file_path: str, content: str):
        if Document is None:
            raise RuntimeError("未安装python-docx，执行 pip install python-docx")
        doc = Document()
        blocks = content.split("\n\n")
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            doc.add_paragraph(block)
        doc.save(file_path)

# ===================== DelTA 多级记忆模块 =====================
class DelTAMemory:
    def __init__(self):
        self.proper_nouns = dict()
        self.bilingual_summary = ""
        self.long_term_context = ""
        self.short_term_context = []
        self.max_short_len = 6

    def update_short_memory(self, target_sent: str):
        self.short_term_context.append(target_sent)
        if len(self.short_term_context) > self.max_short_len:
            self.short_term_context.pop(0)

    def extract_candidate_nouns(self, source_text: str) -> dict:
        clean_text = re.sub(r"\$\$[\s\S]*?\$\$|\$[^\$]*\$", "", source_text)
        word_list = clean_text.split()
        candidate = {}
        for word in word_list:
            if len(word) > 1 and word.istitle() and word.isalpha():
                candidate[word] = ""
        return candidate

    def render_memory_prompt(self) -> str:
        noun_str = "\n".join([f"{k}: {v}" for k, v in self.proper_nouns.items()]) if self.proper_nouns else "无"
        short_str = "\n".join(self.short_term_context) if self.short_term_context else "无"
        memory_block = (
            f"===== 翻译记忆约束（仅约束术语统一，不强制锁死语序）=====\n"
            f"1. 专有名词对照表：\n{noun_str}\n"
            f"2. 全文精简双语摘要：{self.bilingual_summary}\n"
            f"3. 最近6句译文（仅统一术语，句式可灵活调整）：\n{short_str}\n"
            f"=======================================================\n"
        )
        return memory_block

# ===================== 平衡型翻译智能体（新增PDF支持） =====================
class DeltaTranslatorAgent:
    def __init__(self, api_key: Optional[str] = None, model: str = "MiniMax-M3"):
        self.api_key = "sk-cp-5zn8r7SbNKzrr-TBxjv7OxTEXbyYCDRucTg4rBzLco0g1fQVAh_wZ7vVPDRCDxfUt54Ey3MFGHgf5YSeVfLySvCEd2c4exoiQOYorG8CQnAiPaM9KACFJ5g"
        if not self.api_key:
            raise ValueError("需要传入API Key，参数--api-key 或环境变量 ANTHROPIC_API_KEY")
        self.client = anthropic.Anthropic(
            api_key=self.api_key,
            base_url="https://api.minimaxi.com/anthropic"
        )
        self.model = model
        self.max_retry = 10
        self.base_sleep = 6
        self.memory = DelTAMemory()
        self.supported_languages = {
            "zh": "中文", "en": "英文", "ja": "日文", "ko": "韩文",
            "fr": "法文", "de": "德文", "es": "西班牙文", "it": "意大利文",
            "ru": "俄文", "ar": "阿拉伯文", "pt": "葡萄牙文", "auto": "自动检测"
        }
        print("【DelTA翻译智能体｜PDF/md/docx/txt全格式｜公式保护｜学术优化】")

    def _single_chunk_translate(self, chunk_text: str, source_name: str, target_name: str, style_instr: str) -> str:
        memory_info = self.memory.render_memory_prompt()
        prompt = f"""你是专攻NLP/AI计算机学术论文的中英翻译专家，均衡BLEU、CHRF、BERTScore三项指标：
1. BLEU：全文专有名词、专业术语翻译完全统一；
2. BERTScore：译文贴合原文深层语义，中文学术语句通顺自然；
3. CHRF：长句合理拆分，符合国内顶会中文论文表达习惯。
翻译风格：{style_instr}
{memory_info}
# 【绝对强制不可违反规则，违规会造成论文公式/算法损坏】
1. 专有名词、模型缩写首次出现必须标注完整中文译名，如DocMT(文档级机器翻译)、LLM(大语言模型)；
2. 所有单行LaTeX公式 $...$、多行块公式 $$...$$ 完整原样保留，**禁止修改、翻译、删减任何公式内字符、符号、下标**；
3. Markdown ``` 包裹代码块、伪代码、算法伪代码、Python代码完全原样保留，仅翻译代码块外部自然文字；
4. Markdown语法全部保留：#标题、##二级标题、-列表、>引用、|表格、[文本](链接)，仅翻译文本内容，链接URL不动；
5. 论文数学符号、算法变量名、下标L_、$D_s$、$T_1(p)$ 一律不翻译、不改写；
6. Transformer、Qwen、GPT、COMET、LTCR、DocMT、BERT、CNN等NLP标准术语采用国内ICLR/ACL顶会通用中文译法；
7. 术语严格遵循对照表，允许调整语序，但不能更换专有名词译法；
8. 引用标注 (作者, 年份) 完整保留，不改动；
9. 仅输出纯净译文，不要额外解释、注释、说明文字；
10 算法伪代码、for/循环数学流程完全原样输出，禁止删减任意一行代码。

待翻译{source_name}论文片段：
{chunk_text}
"""
        trans_result = ""
        for retry_times in range(self.max_retry):
            try:
                with self.client.messages.stream(
                    model=self.model,
                    max_tokens=100000,
                    messages=[{"role": "user", "content": prompt}]
                ) as stream:
                    trans_result = stream.get_final_message().content[0].text
                time.sleep(random.uniform(1.8, 3.2))
                break
            except RateLimitError as e:
                wait_sec = self.base_sleep * (2 ** retry_times)
                print(f"【限流重试 {retry_times+1}/{self.max_retry}】等待 {wait_sec:.1f}s : {e}")
                time.sleep(wait_sec)
            except APIStatusError as e:
                if 500 <= e.status_code < 600:
                    wait_sec = self.base_sleep * (2 ** retry_times)
                    print(f"【服务异常重试】等待 {wait_sec:.1f}s")
                    time.sleep(wait_sec)
                else:
                    raise
        else:
            raise RuntimeError(f"连续 {self.max_retry} 次API请求失败，当前文本翻译中断")
        return trans_result

    def translate_text(
        self,
        text: str,
        source_lang: str = "auto",
        target_lang: str = "zh",
        style: str = "general"
    ) -> str:
        source_name = self.supported_languages.get(source_lang, source_lang)
        target_name = self.supported_languages.get(target_lang)
        noun_candidate = self.memory.extract_candidate_nouns(text)
        self.memory.proper_nouns.update(noun_candidate)
        style_map = {
            "general": "通用自然书面语，通顺流畅",
            "formal": "正式学术书面文体，严谨规范",
            "casual": "轻松口语化表达，简洁通俗",
            "technical": "AI/NLP计算机学术论文专用，术语标准、公式完整保留、句式符合顶会中文规范"
        }
        style_instr = style_map.get(style, style_map["general"])
        text_chunks = split_long_text(text, max_chunk_chars=1200)
        full_trans = []
        print(f"文本自动拆分为 {len(text)} 段分段翻译，规避token超限截断")
        for idx, chunk in enumerate(text_chunks, 1):
            print(f"正在翻译第 {idx}/{len(text_chunks)} 片段...")
            chunk_trans = self._single_chunk_translate(chunk, source_name, target_name, style_instr)
            full_trans.append(chunk_trans)
            clean_chunk = re.sub(r"\$\$[\s\S]*?\$\$|\$[^\$]*\$", "", chunk)
            for en_word in noun_candidate.keys():
                if en_word in clean_chunk:
                    chinese_tokens = re.findall(r"[\u4e00-\u9fa5]+", chunk_trans)
                    if chinese_tokens:
                        self.memory.proper_nouns[en_word] = chinese_tokens[0]
                    else:
                        self.memory.proper_nouns[en_word] = chunk_trans
        full_result = "\n\n".join(full_trans)
        self.memory.update_short_memory(full_result[:1000])
        if len(self.memory.bilingual_summary) < 600:
            self.memory.bilingual_summary += f"{text[:30]}｜{full_result[:30]}；"
        return full_result

    def translate_file(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        source_lang: str = "auto",
        target_lang: str = "zh",
        style: str = "general",
        pdf_output_type: str = "md"
    ) -> str:
        input_file = Path(input_path)
        if not input_file.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")
        suffix = input_file.suffix.lower()
        temp_dir = tempfile.mkdtemp()
        try:
            # PDF 特殊处理：先转临时md
            if suffix == ".pdf":
                print(f"解析PDF，临时目录：{temp_dir}")
                temp_md = pdf_to_temp_md(str(input_file), temp_dir)
                content = Path(temp_md).read_text(encoding="utf-8")
                # 输出后缀由--pdf-output控制
                out_suffix = f".{pdf_output_type}"
                if output_path is None:
                    out_name = f"{input_file.stem}_delta.translated{out_suffix}"
                    output_path = str(input_file.parent / out_name)
            elif suffix == ".docx":
                if Document is None:
                    raise RuntimeError("翻译docx需安装python-doc: pip install python-docx")
                content = DocxHandler.read_docx(str(input_file))
            elif suffix in [".md", ".txt"]:
                content = input_file.read_text(encoding="utf-8")
            else:
                raise ValueError(f"不支持格式 {suffix}，仅支持 .pdf / .md / .txt / .docx")
            if not content.strip():
                raise ValueError("文件内容为空，无法翻译")
            trans_text = self.translate_text(content, source_lang, target_lang, style)
            out_file = Path(output_path)
            out_suffix = out_file.suffix.lower()
            if out_suffix == ".docx":
                DocxHandler.write_doc(str(out_file), trans_text)
            else:
                out_file.write_text(trans_text, encoding="utf-8")
            return output_path
        finally:
            # 清理临时PDF中转文件夹
            shutil.rmtree(temp_dir, ignore_errors=True)

    def translate_batch(
        self,
        input_paths: list[str],
        output_dir: Optional[str] = None,
        source_lang: str = "auto",
        target_lang: str = "zh",
        style: str = "general",
        max_workers: int = 1,
        pdf_output_type: str = "md"
    ) -> dict:
        import concurrent.futures
        results = {"success": [], "failed": []}
        def single_file_translate(file_path: str):
            try:
                out_path = None
                if output_dir:
                    fp = Path(file_path)
                    suffix = fp.suffix.lower()
                    if suffix == ".pdf":
                        out_suffix = f".{pdf_output_type}"
                    else:
                        out_suffix = suffix
                    out_name = f"{fp.stem}_delta.translated{out_suffix}"
                    out_path = str(Path(output_dir) / out_name)
                res = self.translate_file(file_path, out_path, source_lang, target_lang, style, pdf_output_type)
                return (file_path, res)
            except Exception as err:
                return (file_path, str(err))
        task_list = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            for p in input_paths:
                task = executor.submit(single_file_translate, p)
                task_list.append((task, p))
            for task, file_p in task_list:
                src_path, ret = task.result()
                if isinstance(ret, str) and not ret.startswith("文件不存在"):
                    results["success"].append({"input": src_path, "output": ret})
                    print(f"✓ {src_path} -> {ret}")
                else:
                    results["failed"].append({"input": src_path, "error": ret})
                    print(f"✗ {src_path} 失败: {ret}")
        return results

    def interactive(self, source_lang: str = "auto", target_lang: str = "zh"):
        print("=" * 65)
        print("DelTA交互式翻译｜学术论文优化版｜输入quit退出")
        print("=" * 65)
        while True:
            try:
                user_input = input("> ").strip()
                if user_input.lower() in ("quit", "q", "exit"):
                    print("会话结束，记忆缓存已保存")
                    break
                if not user_input:
                    continue
                res = self.translate_text(user_input, source_lang, target_lang, style="technical")
                print(f"译文：{res}\n")
            except KeyboardInterrupt:
                break

    def evaluate_dataset(self, json_path: str, style: str = "technical", limit: int = None) -> Dict:
        json_file = Path(json_path)
        cache_name = json_file.stem + "_delta_cache.json"
        cache_path = Path(cache_name)
        with open(json_file, "r", encoding="utf-8") as f:
            full_dataset = json.load(f)
        if limit is not None and limit > 0:
            eval_data = full_dataset[:limit]
            print(f"\n限定前 {limit} 条，总{len(full_dataset)}条")
        else:
            eval_data = full_dataset
        valid_src_set = {item["src"].strip() for item in eval_data}
        cache_map = {}
        if cache_path.exists():
            with open(cache_path, "r", encoding="utf-8") as f:
                all_cache = json.load(f)
            filter_cache = [x for x in all_cache if x["source"] in valid_src_set]
            cache_map = {item["source"]: item for item in filter_cache}
            print(f"加载缓存 {cache_name}：{len(cache_map)} 条")
        cache_records = list(cache_map.values())
        predictions = []
        references = []
        new_translate_cnt = 0
        try:
            for idx, item in enumerate(eval_data, 1):
                src_text = item["src"].strip()
                ref_text = item["ref"].strip()
                if src_text in cache_map:
                    record = cache_map[src_text]
                    predictions.append(record["prediction"])
                    references.append([ref_text])
                    continue
                new_translate_cnt += 1
                print(f"\n[{idx}/{len(eval_data)}] 翻译样本：{src_text[:60]}...")
                pred_zh = self.translate_text(src_text, source_lang="en", target_lang="zh", style=style)
                record = {"source": src_text, "reference": ref_text, "prediction": pred_zh}
                cache_records.append(record)
                cache_map[src_text] = record
                with open(cache_path, "w", encoding="utf-8") as f:
                    json.dump(cache_records, f, ensure_ascii=False, indent=2)
                predictions.append(pred_zh)
        except KeyboardInterrupt:
            print("\n手动中断，缓存已保存")
        bleu_corpus = sacrebleu.corpus_bleu(predictions, references)
        bleu_score = bleu_corpus.score
        chrf_corpus = sacrebleu.corpus_chrf(predictions)
        chrf_score = chrf_corpus
        P, R, F1_tensor = score(predictions, references, model_type="bert-base-chinese", lang="zh", verbose=False)
        bert_f1 = float(F1_tensor.mean())
        out_json_name = json_file.stem + "_delta_full_predict.json"
        with open(out_json_name, "w", encoding="utf-8") as f:
            json.dump(cache_records, f, ensure_ascii=False, indent=2)
        return {
            "total_samples": len(eval_data),
            "new_translated_count": new_translate_cnt,
            "BLEU": bleu_score,
            "CHRF": chrf_score,
            "BERTScore_F1": bert_f1,
            "delta_cache_file": str(cache_path),
            "delta_result_file": out_json_name
        }

def main():
    parser = argparse.ArgumentParser(
        description="DelTA翻译智能体【PDF/md/docx/txt四格式】自动识别PDF公式，学术专用",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
# 翻译PDF输出md（推荐）
python translator_delta.py file paper.pdf -s en --style technical --pdf-output md
# 翻译PDF输出docx
python translator_delta.py file paper.pdf -s en --style technical --pdf-output docx
# 翻译md
python translator_delta.py file paper.md -s en --style technical
# 批量PDF全部导出md
python translator_delta batch *.pdf --style technical --pdf-output md -w 1
"""
    )
    subparsers = parser.add_subparsers(dest="command", required=True)
    subparsers.add_parser("interactive", aliases=["i"], help="交互式翻译")
    text_parser = subparsers.add_parser("text", help="翻译单段文本")
    text_parser.add_argument("content")
    text_parser.add_argument("-s", default="auto")
    text_parser.add_argument("-t", default="zh")
    text_parser.add_argument("--style", default="general")
    file_parser = subparsers.add_parser("file", help="翻译单文件 pdf/md/txt/docx")
    file_parser.add_argument("input")
    file_parser.add_argument("-o")
    file_parser.add_argument("-s", default="auto")
    file_parser.add_argument("-t", default="zh")
    file_parser.add_argument("--style", default="general")
    file_parser.add_argument("--pdf-output", choices=["md", "docx"], default="md", help="PDF翻译后输出格式")
    batch_parser = subparsers.add_parser("batch", help="批量翻译文件")
    batch_parser.add_argument("inputs", nargs="+")
    batch_parser.add_argument("-o", "--output-dir")
    batch_parser.add_argument("-s", default="auto")
    batch_parser.add_argument("-t", default="zh")
    batch_parser.add_argument("--style", default="general")
    batch_parser.add_argument("-w", type=int, default=1)
    batch_parser.add_argument("--pdf-output", choices=["md", "docx"], default="md")
    eval_parser = subparsers.add_parser("eval", help="数据集评测")
    eval_parser.add_argument("--json", required=True)
    eval_parser.add_argument("--style", default="technical")
    eval_parser.add_argument("--limit", type=int)
    parser.add_argument("--api-key")
    parser.add_argument("--model", default="MiniMax-M3")
    parser.add_argument("--list-langs", action="store_true")
    args = parser.parse_args()
    agent = DeltaTranslatorAgent(api_key=args.api_key, model=args.model)
    if args.list_langs:
        for k, v in agent.supported_languages.items():
            print(f"{k}: {v}")
        return
    if args.command in ("interactive", "i"):
        agent.interactive()
    elif args.command == "text":
        print(agent.translate_text(args.content, args.s, args.t, args.style))
    elif args.command == "file":
        out = agent.translate_file(
            args.input, args.o, args.s, args.t, args.style, args.pdf_output
        )
        print("输出文件：", out)
    elif args.command == "batch":
        res = agent.translate_batch(
            args.inputs, args.output_dir, args.s, args.t, args.style, args.w, args.pdf_output
        )
        print(f"批量完成：成功{len(res['success'])} 失败{len(res['failed'])}")
    elif args.command == "eval":
        res = agent.evaluate_dataset(args.json, args.style, args.limit)
        print("\n====评测汇总====")
        print(f"总样本：{res['total_samples']}")
        print(f"BLEU {res['BLEU']:.2f} CHRF {res['CHRF']:.2f} BERTScore {res['BERTScore_F1']:.3f}")

if __name__ == "__main__":
    main()